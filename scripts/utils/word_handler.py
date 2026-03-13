"""Word テンプレート処理の共通ユーティリティ

【最重要: run分割問題への対処】
Word は「{{company_name}}」のようなプレースホルダーを
入力・変換の都合で複数の run に分割して保存する。
例: run[0]="{{", run[1]="company", run[2]="_name", run[3]="}}"
この場合、個々の run.text を検索しても placeholder が見つからない。

→ _replace_in_paragraph() では段落内の全 run を結合した全文字列に対して
  置換を行い、結果を run[0] に書き戻す方式を採用。
  これにより run 分割の影響を受けずに確実に置換できる。

【テーブルセル書き込み】
.text への直接代入はセル内の書式（フォント・罫線等）を破壊する。
→ fill_table_cell_safe() では段落の run を書き換える方式を採用。
"""

import re
import sys
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell


# ---------------------------------------------------------------------------
# 内部ユーティリティ
# ---------------------------------------------------------------------------

def _iter_paragraphs(doc: Document) -> Iterator:
    """本文段落とすべてのテーブルセル内段落を横断的に返す"""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def _replace_in_paragraph(para, mapping: dict) -> list[str]:
    """
    段落内のプレースホルダーを置換する。run分割に対応。

    処理方式:
      1. 全 run.text を結合して段落の完全テキストを再現
      2. 全 placeholder を一括置換
      3. 変化があった場合のみ run[0] に書き戻し、残りの run を空文字に
      → run[0] の書式（フォント・サイズ等）を保持する
    """
    if not para.runs:
        return []

    full_text = "".join(run.text for run in para.runs)
    if not any(ph in full_text for ph in mapping):
        return []  # 高速パス

    new_text = full_text
    replaced = []
    for ph, val in mapping.items():
        if ph in new_text:
            new_text = new_text.replace(ph, str(val) if val is not None else "")
            replaced.append(ph)

    if new_text == full_text:
        return []

    # run[0] に新テキストを書き込み、他の run をクリア
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""

    return replaced


def _iter_all_paragraphs(doc: Document) -> Iterator:
    """本文・テーブルセル・ネストテーブルを含む全段落を返す"""
    def _extract(container):
        for para in container.paragraphs:
            yield para
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from _extract(cell)

    yield from _extract(doc)


# ---------------------------------------------------------------------------
# 公開 API
# ---------------------------------------------------------------------------

def load_template(template_path: str) -> Document:
    """テンプレートWordファイルを読み込む"""
    path = Path(template_path)
    if not path.exists():
        print(f"エラー: テンプレートファイルが見つかりません: {template_path}", file=sys.stderr)
        sys.exit(1)
    return Document(template_path)


def replace_placeholders(doc: Document, mappings: dict) -> dict:
    """
    複数のプレースホルダーを一括置換する（run分割対応版）。

    mappings: {"{{key}}": "value", ...}
    戻り値: {"{{key}}": 置換回数, ...}  (0 = 未発見)
    """
    counts = {ph: 0 for ph in mappings}
    active = {ph: val for ph, val in mappings.items() if val is not None and val != ""}

    for para in _iter_all_paragraphs(doc):
        replaced = _replace_in_paragraph(para, active)
        for ph in replaced:
            counts[ph] += 1

    # 未置換のプレースホルダーを警告
    for ph, cnt in counts.items():
        if cnt == 0 and mappings.get(ph):
            print(f"⚠️ 未置換: {ph}  → テンプレート内に存在しない可能性があります", file=sys.stderr)

    return counts


def replace_placeholder(doc: Document, placeholder: str, value: str):
    """単一プレースホルダーの置換（後方互換用）"""
    replace_placeholders(doc, {placeholder: value})


def fill_table_cell_safe(doc: Document, table_index: int, row: int, col: int, value: str):
    """
    指定テーブルのセルに値を書き込む（書式保持版）。

    .text への直接代入は書式を破壊するため、
    段落の run を書き換える方式を採用する。
    セルに段落がない場合は新規追加する。
    """
    if table_index >= len(doc.tables):
        print(f"エラー: テーブルインデックス {table_index} が範囲外です", file=sys.stderr)
        return False
    table = doc.tables[table_index]
    if row >= len(table.rows):
        print(f"エラー: 行インデックス {row} が範囲外です", file=sys.stderr)
        return False
    cells = table.rows[row].cells
    if col >= len(cells):
        print(f"エラー: 列インデックス {col} が範囲外です", file=sys.stderr)
        return False

    cell = cells[col]
    if cell.paragraphs and cell.paragraphs[0].runs:
        # 既存 run に書き込み（書式を保持）
        cell.paragraphs[0].runs[0].text = value
        for run in cell.paragraphs[0].runs[1:]:
            run.text = ""
    elif cell.paragraphs:
        # run がない段落に run を追加
        cell.paragraphs[0].add_run(value)
    else:
        # 段落ごと追加
        cell.add_paragraph(value)
    return True


def fill_table_cell(doc: Document, table_index: int, row: int, col: int, value: str):
    """後方互換用。fill_table_cell_safe を呼び出す"""
    fill_table_cell_safe(doc, table_index, row, col, value)


def save_output(doc: Document, output_path: str):
    """ファイルを保存"""
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"保存完了: {output_path}")


def get_document_structure(doc: Document) -> dict:
    """文書の構造情報を取得"""
    structure = {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "headings": [],
        "table_sizes": [],
    }

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            structure["headings"].append({
                "level": para.style.name,
                "text": para.text,
            })

    for i, table in enumerate(doc.tables):
        structure["table_sizes"].append({
            "index": i,
            "rows": len(table.rows),
            "cols": len(table.rows[0].cells) if table.rows else 0,
        })

    return structure


def list_placeholders(doc: Document) -> list[str]:
    """
    文書内の全プレースホルダー（{{...}}形式）を取得する。
    run分割されたプレースホルダーも段落の完全テキストから検出する。
    """
    pattern = re.compile(r"\{\{.+?\}\}")
    placeholders = set()

    for para in _iter_all_paragraphs(doc):
        # 全 run を結合した完全テキストで検索（run分割対応）
        full_text = "".join(run.text for run in para.runs) if para.runs else para.text
        for match in pattern.finditer(full_text):
            placeholders.add(match.group(0))

    return sorted(placeholders)


def verify_replacement(doc: Document, expected_mappings: dict) -> dict:
    """
    置換が正しく行われたか検証する。

    戻り値: {"{{key}}": {"found": bool, "value": str or None}}
    """
    pattern = re.compile(r"\{\{.+?\}\}")
    remaining = set()

    for para in _iter_all_paragraphs(doc):
        full_text = "".join(run.text for run in para.runs) if para.runs else para.text
        for match in pattern.finditer(full_text):
            remaining.add(match.group(0))

    results = {}
    for ph in expected_mappings:
        results[ph] = {
            "replaced": ph not in remaining,
            "target_value": expected_mappings[ph],
        }
    return results

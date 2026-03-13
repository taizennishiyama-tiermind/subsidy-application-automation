#!/usr/bin/env python3
"""
Word Template Filler

マッピング定義に従って、抽出データをWordテンプレートに埋め込みます。
プレースホルダー置換、動的表生成、フォーマット保持に対応。
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import sys
import re
from pathlib import Path
from copy import deepcopy


class WordTemplateFiller:
    def __init__(self, template_path, mapping_def_path, extracted_data_path):
        """
        Args:
            template_path: Wordテンプレートのパス
            mapping_def_path: マッピング定義JSONのパス
            extracted_data_path: 抽出データJSONのパス
        """
        self.template_path = template_path
        
        # テンプレートを読み込み
        self.doc = Document(template_path)
        
        # マッピング定義を読み込み
        with open(mapping_def_path, 'r', encoding='utf-8') as f:
            self.mapping_def = json.load(f)
        
        # 抽出データを読み込み
        with open(extracted_data_path, 'r', encoding='utf-8') as f:
            self.extracted_data = json.load(f)
        
        self.fill_log = []
    
    def fill(self):
        """データ埋め込みを実行"""
        print("📝 Wordテンプレートへのデータ埋め込み開始...")
        print("=" * 60)
        
        # 1. 段落内のプレースホルダー置換
        if "placeholder_mappings" in self.mapping_def:
            self._replace_placeholders_in_paragraphs()
        
        # 2. 表内のプレースホルダー置換 & 動的表生成
        if "table_mappings" in self.mapping_def:
            self._fill_tables()
        
        print("\n" + "=" * 60)
        print("✅ 埋め込み完了")
    
    def _get_value_from_path(self, data_path):
        """data_path (例: "company_info.name") から値を取得"""
        keys = data_path.split(".")
        value = self.extracted_data
        
        for key in keys:
            if isinstance(value, dict):
                value = value.get(key)
            else:
                return None
        
        return value
    
    def _replace_placeholders_in_paragraphs(self):
        """段落内のプレースホルダーを置換"""
        print("\n[段落内プレースホルダー置換]")
        
        placeholder_map = {}
        for mapping in self.mapping_def.get("placeholder_mappings", []):
            placeholder = mapping["placeholder"]
            data_path = mapping["data_path"]
            value = self._get_value_from_path(data_path)
            
            if value is not None:
                placeholder_map[placeholder] = str(value)
                print(f"  {placeholder} → {value}")
            else:
                print(f"  ⚠️ {placeholder}: データなし")
        
        # すべての段落を走査
        for paragraph in self.doc.paragraphs:
            original_text = paragraph.text
            
            # プレースホルダーがあるか確認
            modified = False
            for placeholder, value in placeholder_map.items():
                if placeholder in original_text:
                    modified = True
                    break
            
            if modified:
                # ランを保持したまま置換
                self._replace_in_paragraph(paragraph, placeholder_map)
    
    def _replace_in_paragraph(self, paragraph, placeholder_map):
        """
        段落内のプレースホルダーを置換
        
        重要: ランのフォーマットを保持する
        """
        # 段落全体のテキストを取得
        full_text = paragraph.text
        
        # 置換を実行
        for placeholder, value in placeholder_map.items():
            full_text = full_text.replace(placeholder, value)
        
        # ランをクリア
        for run in paragraph.runs:
            run.text = ""
        
        # 最初のランに新しいテキストを設定(フォーマット保持)
        if paragraph.runs:
            paragraph.runs[0].text = full_text
        else:
            paragraph.add_run(full_text)
    
    def _fill_tables(self):
        """表の処理"""
        print("\n[表の処理]")
        
        for table_mapping in self.mapping_def.get("table_mappings", []):
            table_index = table_mapping.get("table_index", 0)
            
            if table_index >= len(self.doc.tables):
                print(f"  ⚠️ 表{table_index}が見つかりません")
                continue
            
            table = self.doc.tables[table_index]
            
            if table_mapping.get("is_dynamic", False):
                # 動的表: 行を繰り返す
                self._fill_dynamic_table(table, table_mapping)
            else:
                # 静的表: プレースホルダー置換のみ
                self._fill_static_table(table, table_mapping)
    
    def _fill_static_table(self, table, table_mapping):
        """静的表のプレースホルダー置換"""
        print(f"  静的表(表{table_mapping.get('table_index', 0)})")
        
        placeholder_map = {}
        for cell_mapping in table_mapping.get("cell_mappings", []):
            placeholder = cell_mapping["placeholder"]
            data_path = cell_mapping["data_path"]
            value = self._get_value_from_path(data_path)
            
            if value is not None:
                placeholder_map[placeholder] = str(value)
        
        # すべてのセルを走査して置換
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    
                    for placeholder, value in placeholder_map.items():
                        if placeholder in original_text:
                            self._replace_in_paragraph(paragraph, {placeholder: value})
                            print(f"    {placeholder} → {value}")
    
    def _fill_dynamic_table(self, table, table_mapping):
        """
        動的表: テンプレート行を繰り返す
        
        通常、2行目がテンプレート行(プレースホルダーを含む)
        """
        print(f"  動的表(表{table_mapping.get('table_index', 0)})")
        
        data_path = table_mapping["data_path"]
        table_data = self._get_value_from_path(data_path)
        
        if not isinstance(table_data, list):
            print(f"    ⚠️ {data_path} はリストではありません")
            return
        
        if len(table_data) == 0:
            print(f"    ⚠️ データが空です")
            return
        
        # テンプレート行のインデックス(通常は1 = 2行目)
        template_row_index = table_mapping.get("template_row_index", 1)
        
        if template_row_index >= len(table.rows):
            print(f"    ⚠️ テンプレート行{template_row_index}が存在しません")
            return
        
        template_row = table.rows[template_row_index]
        
        # 列マッピング
        column_mappings = table_mapping.get("column_mappings", [])
        
        # 既存のデータ行を削除(テンプレート行以降)
        rows_to_delete = len(table.rows) - template_row_index - 1
        for _ in range(rows_to_delete):
            table._element.remove(table.rows[-1]._element)
        
        # データ行を追加
        for item in table_data:
            # 新しい行を追加(テンプレート行をコピー)
            new_row = self._copy_table_row(table, template_row)
            
            # 各セルにデータを埋め込む
            for col_idx, col_mapping in enumerate(column_mappings):
                if col_idx >= len(new_row.cells):
                    continue
                
                data_field = col_mapping["data_field"]
                value = item.get(data_field, "")
                
                # セルのテキストを置換
                cell = new_row.cells[col_idx]
                for paragraph in cell.paragraphs:
                    paragraph.text = str(value)
            
            print(f"    ✓ 行追加: {item.get(column_mappings[0]['data_field'], '')}")
        
        # テンプレート行を削除
        table._element.remove(template_row._element)
    
    def _copy_table_row(self, table, row):
        """
        表の行をコピーして新しい行を追加
        
        フォーマットを保持
        """
        # 新しい行を追加
        new_row = table.add_row()
        
        # 各セルのスタイルをコピー
        for idx, cell in enumerate(row.cells):
            if idx < len(new_row.cells):
                new_cell = new_row.cells[idx]
                
                # セルの段落スタイルをコピー
                for src_para, dst_para in zip(cell.paragraphs, new_cell.paragraphs):
                    dst_para.alignment = src_para.alignment
                    
                    # ランのフォーマットをコピー
                    if src_para.runs:
                        dst_para.runs[0].bold = src_para.runs[0].bold
                        dst_para.runs[0].italic = src_para.runs[0].italic
                        dst_para.runs[0].font.size = src_para.runs[0].font.size
                        dst_para.runs[0].font.name = src_para.runs[0].font.name
        
        return new_row
    
    def save(self, output_path):
        """完成ファイルを保存"""
        self.doc.save(output_path)
        print(f"\n💾 保存完了: {output_path}")
    
    def verify(self):
        """埋め込み結果を検証"""
        print("\n" + "=" * 60)
        print("🔍 検証中...")
        print("=" * 60)
        
        # プレースホルダーが残っていないかチェック
        remaining_placeholders = []
        
        for paragraph in self.doc.paragraphs:
            if re.search(r'\{\{|\{[A-Za-z]|\[', paragraph.text):
                remaining_placeholders.append(paragraph.text[:100])
        
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if re.search(r'\{\{|\{[A-Za-z]|\[', paragraph.text):
                            remaining_placeholders.append(paragraph.text[:100])
        
        if remaining_placeholders:
            print("⚠️ 未置換のプレースホルダーが残っています:")
            for ph in remaining_placeholders[:5]:
                print(f"  {ph}")
            return False
        else:
            print("✅ 検証OK - すべてのプレースホルダーが置換されました")
            return True


def main():
    if len(sys.argv) < 4:
        print("使い方: python fill_word_template.py <テンプレート.docx> <マッピング定義.json> <抽出データ.json> [出力先.docx]")
        print("\n例:")
        print("  python fill_word_template.py template.docx word_mapping.json extracted_data.json output.docx")
        sys.exit(1)
    
    template_path = sys.argv[1]
    mapping_path = sys.argv[2]
    data_path = sys.argv[3]
    output_path = sys.argv[4] if len(sys.argv) > 4 else "completed_application.docx"
    
    # ファイル存在チェック
    for path, name in [(template_path, "テンプレート"), (mapping_path, "マッピング定義"), (data_path, "抽出データ")]:
        if not Path(path).exists():
            print(f"❌ エラー: {name}ファイルが見つかりません: {path}")
            sys.exit(1)
    
    # 埋め込み実行
    filler = WordTemplateFiller(template_path, mapping_path, data_path)
    filler.fill()
    
    # 検証
    is_valid = filler.verify()
    
    # 保存
    filler.save(output_path)
    
    if not is_valid:
        print("\n⚠️ 検証で問題が見つかりましたが、ファイルは保存されました。")
        print("   手動で確認してください。")
    else:
        print("\n🎉 完成! 申請書が正常に生成されました。")


if __name__ == "__main__":
    main()
